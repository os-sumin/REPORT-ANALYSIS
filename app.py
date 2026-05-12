import streamlit as st
import pandas as pd
import pdfplumber
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os
from datetime import datetime
from openai import OpenAI

# 페이지 설정
st.set_page_config(
    page_title="사업화분석 보고서 자동 생성",
    page_icon="📊",
    layout="wide"
)

# OpenAI API 키 설정 (환경 변수 또는 직접 입력)
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")

# 제목
st.title("📊 사업화분석 보고서 자동 생성 시스템")
st.markdown("---")

# 사이드바
with st.sidebar:
    st.header("📋 사용 가이드")
    st.markdown("""
    ### 1️⃣ 파일 준비
    - 최종보고서 PDF (선택)
    - 과제정보 엑셀 (필수)
    - 재무정보 엑셀 (선택)
    
    ### 2️⃣ GPT API 설정
    - OpenAI API 키 입력
    
    ### 3️⃣ 파일 업로드
    - 최소 과제정보 엑셀 필요
    
    ### 4️⃣ 보고서 생성
    - '보고서 생성' 버튼 클릭
    - 30~60초 대기
    
    ### 5️⃣ 다운로드
    - 생성된 DOCX 다운로드
    """)
    
    st.markdown("---")
    
    # API 키 입력
    st.subheader("🔑 OpenAI API 키")
    api_key_input = st.text_input(
        "API 키 입력",
        value=OPENAI_API_KEY,
        type="password",
        help="https://platform.openai.com/api-keys 에서 발급"
    )
    
    if api_key_input:
        _trimmed = api_key_input.strip()
        try:
            _trimmed.encode('ascii')
            _is_ascii = True
        except UnicodeEncodeError:
            _is_ascii = False
        if not _is_ascii:
            st.error(
                "❌ API 키에 한글/특수문자가 섞여있습니다. "
                "한글 입력기를 끄고 영문/숫자만 붙여넣어 주세요."
            )
        elif not _trimmed.startswith('sk-'):
            st.warning("⚠️ API 키는 'sk-' 로 시작해야 합니다. 형식을 확인해주세요.")
        else:
            st.session_state["openai_api_key"] = _trimmed
            st.success(f"✅ API 키 설정됨 (길이 {len(_trimmed)})")
    else:
        st.warning("⚠️ GPT 분석을 위해 API 키가 필요합니다")
    
    st.markdown("---")
    st.markdown("### 📥 템플릿 다운로드")
    
    # 템플릿 다운로드 버튼
    try:
        with open("과제정보_템플릿.xlsx", "rb") as f:
            st.download_button(
                label="📊 과제정보 템플릿",
                data=f,
                file_name="과제정보_템플릿.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
        
        with open("재무정보_템플릿.xlsx", "rb") as f:
            st.download_button(
                label="💰 재무정보 템플릿",
                data=f,
                file_name="재무정보_템플릿.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
    except:
        st.info("템플릿은 첫 실행 시 자동 생성됩니다")

# 메인 영역
col1, col2, col3 = st.columns(3)

with col1:
    st.subheader("📄 최종보고서 PDF (선택)")
    pdf_file = st.file_uploader("PDF 파일 업로드", type=['pdf'], key="pdf")
    if pdf_file:
        st.success(f"✅ {pdf_file.name}")

with col2:
    st.subheader("📊 과제정보 엑셀 (필수)")
    project_file = st.file_uploader("과제정보 엑셀 업로드", type=['xlsx', 'xls'], key="project")
    if project_file:
        st.success(f"✅ {project_file.name}")

with col3:
    st.subheader("💰 재무정보 엑셀 (선택)")
    finance_file = st.file_uploader("재무정보 엑셀 업로드", type=['xlsx', 'xls'], key="finance")
    if finance_file:
        st.success(f"✅ {finance_file.name}")

st.markdown("---")

# 파일 파싱 함수
def parse_pdf(pdf_file):
    """PDF에서 텍스트 추출 (페이지 구분 보존)"""
    try:
        with pdfplumber.open(pdf_file) as pdf:
            pages = []
            for page in pdf.pages:
                # 한국어 세로조판/사이드바를 어느 정도 완화하기 위해 tolerance 조정
                t = page.extract_text(x_tolerance=2, y_tolerance=3) or ""
                if t.strip():
                    pages.append(t)
        return "\n\n".join(pages)
    except Exception as e:
        return f"PDF 파싱 실패: {str(e)}"


def _clean_pdf_text(text: str) -> str:
    """추출된 PDF 텍스트의 흔한 잡음 정리.

    - 한 줄에 한 글자만 있는 라인(세로 사이드바)을 합침
    - 과도한 공백/줄바꿈 압축
    """
    import re
    lines = text.split('\n')
    cleaned = []
    buf_single = []
    for line in lines:
        stripped = line.strip()
        if len(stripped) <= 1:
            if stripped:
                buf_single.append(stripped)
            continue
        if buf_single:
            cleaned.append(''.join(buf_single))
            buf_single = []
        cleaned.append(stripped)
    if buf_single:
        cleaned.append(''.join(buf_single))
    out = '\n'.join(cleaned)
    out = re.sub(r'\n{3,}', '\n\n', out)
    out = re.sub(r'[ \t]{2,}', ' ', out)
    return out.strip()

def parse_project_excel(file):
    """과제정보 엑셀 파싱 - 유연한 버전"""
    try:
        # 모든 시트 확인
        xl = pd.ExcelFile(file)
        
        # '과제정보' 시트 찾기
        sheet_name = None
        for name in xl.sheet_names:
            if '과제' in name:
                sheet_name = name
                break
        
        if sheet_name is None:
            sheet_name = xl.sheet_names[0]  # 첫 번째 시트 사용
        
        df = pd.read_excel(file, sheet_name=sheet_name)
        
        # 필수 컬럼 유연하게 체크
        required = ['과제명', '기업명']
        missing = [col for col in required if col not in df.columns]
        
        if missing:
            st.error(f"❌ 필수 컬럼이 없습니다: {', '.join(missing)}")
            st.info(f"현재 컬럼: {', '.join(df.columns)}")
            return None
        
        # 빈 행 제거
        df = df.dropna(subset=['과제명', '기업명'])
        
        # 빈 값 처리
        df = df.fillna('')
        
        return df.to_dict('records')
    except Exception as e:
        st.error(f"과제정보 파싱 오류: {str(e)}")
        return None

def parse_finance_excel(file):
    """재무정보 엑셀 파싱 - 유연한 버전"""
    try:
        # 모든 시트 확인
        xl = pd.ExcelFile(file)
        
        # '재무' 시트 찾기
        sheet_name = None
        for name in xl.sheet_names:
            if '재무' in name:
                sheet_name = name
                break
        
        if sheet_name is None:
            sheet_name = xl.sheet_names[0]  # 첫 번째 시트 사용
        
        df = pd.read_excel(file, sheet_name=sheet_name)
        
        # 필수 컬럼 유연하게 체크
        required = ['연도']
        
        # 과제명 또는 기업명 중 하나라도 있으면 OK
        if '과제명' not in df.columns and '기업명' not in df.columns:
            st.warning("⚠️ '과제명' 또는 '기업명' 컬럼이 필요합니다")
            st.info(f"현재 컬럼: {', '.join(df.columns)}")
            return None
        
        if '연도' not in df.columns:
            st.warning("⚠️ '연도' 컬럼이 필요합니다")
            return None
        
        # 빈 행 제거
        df = df.dropna(subset=['연도'])
        
        # 빈 값 처리
        df = df.fillna('')
        
        return df.to_dict('records')
    except Exception as e:
        st.error(f"재무정보 파싱 오류: {str(e)}")
        return None

def call_gpt_analysis(company_name, project_name, api_key):
    """GPT API 호출 - 시장 분석"""
    if not api_key:
        return {
            'market_size': '[GPT API 키가 필요합니다]',
            'competitors': '[GPT API 키가 필요합니다]',
            'strategy': '[GPT API 키가 필요합니다]'
        }

    # API 키 정리: 양끝 공백/줄바꿈 제거, 비-ASCII 검사
    cleaned_key = api_key.strip()
    try:
        cleaned_key.encode('ascii')
    except UnicodeEncodeError:
        msg = (
            "OpenAI API 키에 한글 또는 보이지 않는 특수문자가 섞여 있습니다. "
            "키를 복사할 때 한글 입력기가 켜져 있거나 메신저를 통해 받을 때 "
            "보이지 않는 문자가 섞이는 경우가 흔합니다. "
            "https://platform.openai.com/api-keys 에서 새로 발급받아 영문/숫자만 붙여넣어 주세요."
        )
        return {
            'market_size': f'GPT 분석 오류: {msg}',
            'competitors': '오류 발생',
            'strategy': '오류 발생'
        }
    if not cleaned_key.startswith('sk-'):
        return {
            'market_size': 'GPT 분석 오류: API 키는 "sk-" 로 시작해야 합니다. 키 형식을 확인해주세요.',
            'competitors': '오류 발생',
            'strategy': '오류 발생'
        }

    try:
        prompt = f"""
다음 기업의 사업화 분석을 해주세요:

기업명: {company_name}
과제명: {project_name}

아래 항목을 분석해주세요:
1. 시장 환경 (시장 규모, 성장률, TAM/SAM/SOM)
2. 주요 경쟁사 3~5개 (국내외 포함)
3. 사업화 전략 제안 (단기/중기/장기)

각 항목을 300자 이내로 간결하게 작성해주세요.
"""

        client = OpenAI(api_key=cleaned_key)
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "당신은 기술 사업화 전문 컨설턴트입니다."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=2000
        )

        result = response.choices[0].message.content
        
        # 간단한 파싱
        parts = result.split('\n\n')
        
        return {
            'market_size': parts[0] if len(parts) > 0 else result,
            'competitors': parts[1] if len(parts) > 1 else '분석 중...',
            'strategy': parts[2] if len(parts) > 2 else '분석 중...'
        }
        
    except Exception as e:
        return {
            'market_size': f'GPT 분석 오류: {str(e)}',
            'competitors': '오류 발생',
            'strategy': '오류 발생'
        }

def generate_report(pdf_text, project_data, finance_data, use_gpt=False, api_key=""):
    """사업화분석 보고서 생성 - GPT 연동 버전"""
    doc = Document()
    
    # 제목
    title = doc.add_heading('사업화분석 보고서', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 부제목
    if project_data:
        company = project_data[0].get('기업명', '')
        project_name = project_data[0].get('과제명', '')
        
        subtitle = doc.add_paragraph()
        subtitle.add_run(f"{company}\n{project_name}").bold = True
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        company = "알 수 없음"
        project_name = "알 수 없음"
    
    doc.add_paragraph()
    
    # 생성일
    doc.add_paragraph(f"생성일: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_page_break()
    
    # 목차
    doc.add_heading('목차', 1)
    toc_items = [
        "1. 과제 개요",
        "2. 기업 정보",
        "3. 기술 성과",
        "4. 재무 현황",
        "5. 시장 환경 분석",
        "6. 경쟁사 분석",
        "7. 사업화 전략 제안"
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # 1. 과제 개요
    doc.add_heading('1. 과제 개요', 1)
    
    if project_data:
        p = project_data[0]
        
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        rows_data = [
            ('과제명', p.get('과제명', '')),
            ('기업명', p.get('기업명', '')),
            ('연구기간', f"{p.get('연구기간_시작', '')} ~ {p.get('연구기간_종료', '')}"),
            ('정부지원금', f"{p.get('정부지원금(원)', ''):,} 원" if p.get('정부지원금(원)') else '정보 없음'),
            ('주관기관', p.get('주관기관', '정보 없음')),
        ]
        
        for i, (label, value) in enumerate(rows_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)
    
    doc.add_page_break()
    
    # GPT 분석 (선택적)
    gpt_analysis = None
    if use_gpt and api_key:
        with st.spinner("🤖 GPT 분석 중..."):
            gpt_analysis = call_gpt_analysis(company, project_name, api_key)
    
    # 2. 기업 정보
    doc.add_heading('2. 기업 정보', 1)
    if gpt_analysis:
        doc.add_paragraph(gpt_analysis.get('market_size', '[분석 예정]'))
    else:
        doc.add_paragraph("[GPT 분석이 활성화되지 않았습니다]")
    
    doc.add_page_break()
    
    # 3. 기술 성과
    doc.add_heading('3. 기술 성과', 1)
    if pdf_text and len(pdf_text) > 100:
        cleaned = _clean_pdf_text(pdf_text)
        for para in cleaned.split('\n\n'):
            para = para.strip()
            if para:
                doc.add_paragraph(para)
    else:
        doc.add_paragraph("[PDF 파일을 업로드하면 자동으로 추출됩니다]")
    
    doc.add_page_break()
    
    # 4. 재무 현황
    doc.add_heading('4. 재무 현황', 1)
    
    if finance_data and len(finance_data) > 0:
        doc.add_heading('연도별 재무 지표', 2)
        
        # 사용 가능한 컬럼 확인
        available_cols = list(finance_data[0].keys())
        display_cols = ['연도']
        
        for col in ['부채총계(원)', '자산총계(원)', '매출액(원)', '영업이익(원)', '당기순이익(원)']:
            if col in available_cols:
                display_cols.append(col)
        
        table = doc.add_table(rows=1, cols=len(display_cols))
        table.style = 'Light Grid Accent 1'
        
        # 헤더
        for i, col in enumerate(display_cols):
            table.rows[0].cells[i].text = col.replace('(원)', '')
        
        # 데이터
        for f in finance_data:
            row = table.add_row()
            for i, col in enumerate(display_cols):
                value = f.get(col, '')
                if col == '연도':
                    if isinstance(value, (int, float)) and value != '':
                        row.cells[i].text = f"{int(value)}"
                    else:
                        row.cells[i].text = str(value) if value else '-'
                elif isinstance(value, (int, float)) and value != '':
                    row.cells[i].text = f"{value:,.0f}"
                else:
                    row.cells[i].text = str(value) if value else '-'
    else:
        doc.add_paragraph("[재무정보 파일을 업로드하면 자동으로 표시됩니다]")
    
    doc.add_page_break()
    
    # 5. 시장 환경 분석
    doc.add_heading('5. 시장 환경 분석', 1)
    if gpt_analysis:
        doc.add_paragraph(gpt_analysis.get('market_size', '[분석 예정]'))
    else:
        doc.add_paragraph("[GPT API를 활성화하면 자동으로 생성됩니다]")
    
    doc.add_page_break()
    
    # 6. 경쟁사 분석
    doc.add_heading('6. 경쟁사 분석', 1)
    if gpt_analysis:
        doc.add_paragraph(gpt_analysis.get('competitors', '[분석 예정]'))
    else:
        doc.add_paragraph("[GPT API를 활성화하면 자동으로 생성됩니다]")
    
    doc.add_page_break()
    
    # 7. 사업화 전략 제안
    doc.add_heading('7. 사업화 전략 제안', 1)
    if gpt_analysis:
        doc.add_paragraph(gpt_analysis.get('strategy', '[분석 예정]'))
    else:
        doc.add_paragraph("[GPT API를 활성화하면 자동으로 생성됩니다]")
    
    # 메모리에 저장
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

# GPT 사용 옵션
use_gpt = st.checkbox(
    "🤖 GPT 분석 사용 (시장 분석, 경쟁사 분석, 전략 제안)",
    value=True,
    help="OpenAI API 키가 필요합니다. 체크 해제 시 기본 템플릿만 생성됩니다."
)

# 생성 버튼
if st.button("🚀 보고서 생성", type="primary", use_container_width=True):
    if not project_file:
        st.error("❌ 최소한 과제정보 엑셀을 업로드해주세요!")
    else:
        try:
            with st.spinner("📊 보고서 생성 중... (30~60초 소요)"):
                # 파일 파싱
                pdf_text = parse_pdf(pdf_file) if pdf_file else ""
                project_data = parse_project_excel(project_file)
                
                if project_data is None:
                    st.error("과제정보 파일을 확인해주세요")
                    st.stop()
                
                finance_data = parse_finance_excel(finance_file) if finance_file else None
                
                # API 키 확인
                api_key = api_key_input or ""

                if use_gpt and not api_key:
                    st.warning("⚠️ GPT 분석을 사용하려면 API 키를 입력해주세요. 기본 템플릿으로 생성합니다.")

                # 보고서 생성
                report_buffer = generate_report(
                    pdf_text,
                    project_data,
                    finance_data,
                    use_gpt=use_gpt,
                    api_key=api_key
                )
                
                # 성공 메시지
                st.success("✅ 보고서 생성 완료!")
                
                # 다운로드 버튼
                company_name = project_data[0].get('기업명', '기업') if project_data else '기업'
                filename = f"사업화분석보고서_{company_name}_{datetime.now().strftime('%Y%m%d')}.docx"
                
                st.download_button(
                    label="📥 보고서 다운로드",
                    data=report_buffer,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                
                # 미리보기
                st.info(f"""
                ### 생성된 보고서 정보
                - 기업명: {project_data[0].get('기업명', '-')}
                - 과제명: {project_data[0].get('과제명', '-')}
                - 재무 데이터: {len(finance_data) if finance_data else 0}개 연도
                - GPT 분석: {'사용' if (use_gpt and api_key) else '미사용'}
                - 생성 시간: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
                """)
                
        except Exception as e:
            st.error(f"❌ 오류 발생: {str(e)}")
            st.exception(e)

# 푸터
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: gray;'>
    💡 GPT API 사용 시 OpenAI 계정의 요금이 부과될 수 있습니다 (보고서당 약 $0.01~0.05)
</div>
""", unsafe_allow_html=True)
